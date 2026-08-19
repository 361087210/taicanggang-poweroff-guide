from docx import Document
from docx.oxml.ns import qn
import os, json, re, shutil

DOCX_PATH = r"c:\Users\36108\.trae-cn\attachments\6a848b954d2a63c19205b4c0\92896d6e-3683-485f-9431-84802b0c9a99_15d970d7-7e16-48ba-afde-fd572facd54c_太仓港商品车断电操作手册20260603版.docx"
IMG_DIR = r"e:\车辆断电指导应用开发\docx_content\images"
DEMO_IMG_DIR = r"e:\车辆断电指导应用开发\vehicle_images"
OUTPUT_DIR = r"e:\车辆断电指导应用开发\docx_content"

os.makedirs(DEMO_IMG_DIR, exist_ok=True)

doc = Document(DOCX_PATH)

image_map = {}
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_map[rel.rId] = os.path.basename(rel.target_ref)

def get_images_from_element(element):
    imgs = []
    for blip in element.iter(qn('a:blip')):
        embed = blip.get(qn('r:embed'))
        if embed and embed in image_map:
            imgs.append(image_map[embed])
    return imgs

# Track heading hierarchy
body = doc.element.body
vehicles = []
heading_stack = {"h1": "", "h2": "", "h3": ""}

brand_map = {
    "比亚迪(BYD)": ("比亚迪", "BYD"),
    "长安（新能源及混动车型无需断电池负极，车内下电后锁门即可）": ("长安", "Changan"),
    "上汽": ("上汽", "SAIC"),
    "长城": ("长城", "GreatWall"),
    "东风": ("东风", "DongFeng"),
    "江淮（JAC）": ("江淮", "JAC"),
    "吉利": ("吉利", "Geely"),
    "奇瑞（Chery）": ("奇瑞", "Chery"),
    "东南SOUEAST": ("东南", "Soueast"),
    "零跑(Leapmotor)": ("零跑", "Leapmotor"),
    "广汽": ("广汽", "GAC")
}

power_type_map = {
    "海豚": "纯电", "海鸥": "纯电", "海豹": "纯电", "海狮": "纯电",
    "元UP": "纯电", "元PRO": "纯电", "元PLUS": "纯电",
    "唐L-EV": "纯电", "唐EV": "纯电", "ATTO8": "纯电",
    "极氪": "纯电", "好猫": "纯电", "EX5": "纯电", "EX2": "纯电",
    "领克02": "混动", "领克08": "混动",
    "AION": "纯电", "零跑B10": "纯电", "零跑C10": "纯电",
}

for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

    if tag == 'p':
        for para in doc.paragraphs:
            if para._element is child:
                text = para.text.strip()
                style_name = para.style.name if para.style else ''

                if text == '太仓港商品车断电操作手册':
                    break

                if 'Heading 1' in style_name:
                    heading_stack["h1"] = text
                    heading_stack["h2"] = ""
                    heading_stack["h3"] = ""
                elif 'Heading 2' in style_name:
                    heading_stack["h2"] = text
                    heading_stack["h3"] = ""
                elif 'Heading 3' in style_name or (text and heading_stack["h2"] and not heading_stack["h3"]):
                    heading_stack["h3"] = text
                break

    elif tag == 'tbl':
        for tbl in doc.tables:
            if tbl._element is child:
                table = []
                for row in tbl.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        cell_images = get_images_from_element(cell._element)
                        row_data.append({"text": cell_text, "images": cell_images})
                    table.append(row_data)

                # Collect all text and images from the table
                all_text = []
                all_images = []
                for row in table:
                    for cell in row:
                        if cell["text"]:
                            all_text.append(cell["text"])
                        all_images.extend(cell["images"])

                # Extract photos (from first 2 columns of rows 1-2)
                photos = []
                for ri in [1, 2]:
                    if ri < len(table):
                        for ci in [0, 1]:
                            if ci < len(table[ri]):
                                for img in table[ri][ci]["images"]:
                                    if img not in photos and img != "image_1.png":
                                        photos.append(img)

                # Copy images to demo directory
                for img in photos:
                    src = os.path.join(IMG_DIR, img)
                    dst = os.path.join(DEMO_IMG_DIR, img)
                    if os.path.exists(src) and not os.path.exists(dst):
                        shutil.copy2(src, dst)

                # Extract key handling
                key_frame = []
                key_container = []
                for row in table:
                    for cell in row:
                        txt = cell["text"]
                        if "所有车辆钥匙" in txt or ("钥匙数量" in txt and "绑扎" in txt):
                            if "铅封袋" in txt or "铅封" in txt:
                                key_frame = [s.strip() for s in txt.split('\n') if s.strip()]
                            elif "进箱" in txt or "中控台" in txt:
                                key_container = [s.strip() for s in txt.split('\n') if s.strip()]

                # Fallback: check specific cells
                if not key_frame and len(table) > 2 and len(table[2]) > 2:
                    txt = table[2][2]["text"]
                    if txt and "钥匙" in txt:
                        key_frame = [s.strip() for s in txt.split('\n') if s.strip()]
                if not key_container and len(table) > 2 and len(table[2]) > 3:
                    txt = table[2][3]["text"]
                    if txt and ("钥匙" in txt or "进箱" in txt or "中控" in txt):
                        key_container = [s.strip() for s in txt.split('\n') if s.strip()]

                # Find position
                position = ""
                for row in table:
                    for cell in row:
                        txt = cell["text"]
                        if "断电位置" in txt:
                            cleaned = txt.replace("断电位置：", "").replace("断电位置:", "").replace("断电位置", "").strip()
                            if cleaned and cleaned != "断电位置":
                                position = cleaned
                                break
                    if position:
                        break

                # If no position found with "断电位置" prefix, look for position-like text in row 4
                if not position and len(table) > 4:
                    txt = table[4][0]["text"] if len(table[4]) > 0 else ""
                    if txt and txt not in ["断电位置", "断电步骤", "断电视频", "备注", "车辆照片", "钥匙处理方式", "框架", "集装箱"]:
                        if not any(h in txt for h in ["断电", "车辆照片", "钥匙", "框架", "集装箱", "备注"]):
                            position = txt

                # Find steps
                steps_text = ""
                for row in table:
                    for ci, cell in enumerate(row):
                        txt = cell["text"]
                        if txt and re.match(r'^1\.', txt) and "短按" in txt or ("扳手" in txt and "步骤" not in txt):
                            steps_text = txt
                            break
                    if steps_text:
                        break
                # Fallback: check column 1 of rows 3-5
                if not steps_text:
                    for ri in range(3, len(table)):
                        if len(table[ri]) > 1:
                            txt = table[ri][1]["text"]
                            if txt and (re.match(r'^1\.', txt) or "短按" in txt or "扳手" in txt):
                                steps_text = txt
                                break

                steps = [s.strip() for s in steps_text.split('\n') if s.strip() and re.match(r'^\d+\.', s.strip())]

                # Find remarks
                remarks = ""
                for ri in range(3, len(table)):
                    if len(table[ri]) > 3:
                        txt = table[ri][3]["text"]
                        if txt and txt not in ["备注", "放干燥剂"] and "钥匙" not in txt and "断电" not in txt and "车辆" not in txt:
                            remarks = txt
                if not remarks:
                    # Check for "放干燥剂" as default remark
                    for row in table:
                        for cell in row:
                            if "放干燥剂" in cell["text"]:
                                remarks = "放干燥剂"
                                break

                # Build vehicle name
                brand_raw = heading_stack["h1"]
                brand_info = brand_map.get(brand_raw, (brand_raw, ""))
                brand_cn = brand_info[0]
                brand_en = brand_info[1] if len(brand_info) > 1 else ""
                series = heading_stack["h2"]
                config = heading_stack["h3"]

                # Display name
                display = f"{brand_cn}{series}"
                if config and config != series:
                    display += f"({config})"

                # Determine power type
                power_type = "混动"
                for key, pt in power_type_map.items():
                    if key in display or key in series:
                        power_type = pt
                        break
                if "EV" in config or "EV" in series:
                    power_type = "纯电"
                if "DM" in config or "DM" in series or "PHEV" in config:
                    power_type = "混动"

                # For "无需断电" vehicles
                if "无需断电" in position or "遥控钥匙锁车" in position:
                    if not steps:
                        steps = ["遥控钥匙锁车，无需断电池负极"]

                vehicle = {
                    "id": len(vehicles) + 1,
                    "brand": brand_cn,
                    "brandEn": brand_en,
                    "series": series,
                    "config": config or "标准",
                    "display": display,
                    "powerType": power_type,
                    "position": position or "前机盖电池负极",
                    "steps": steps if steps else ["打开主驾驶车门，确认全部车窗关闭，取出车钥匙"],
                    "keyFrame": key_frame if key_frame else ["钥匙数量绑扎检查完", "确认断电无误后放到铅封袋内封好袋口", "放入钥匙盒关好，封好铅封"],
                    "keyContainer": key_container if key_container else ["车辆钥匙数量及绑扎检查完", "车辆进箱无需收钥匙，确认断电无误后放置于车内中控台"],
                    "remarks": remarks,
                    "photos": [f"vehicle_images/{img}" for img in photos],
                    "photoCount": len(photos),
                    "videos": 0
                }
                vehicles.append(vehicle)
                break

print(f"Total vehicles: {len(vehicles)}")

# Count by brand
brand_counts = {}
for v in vehicles:
    b = v["brand"]
    brand_counts[b] = brand_counts.get(b, 0) + 1

print("\n品牌分布:")
for b, c in brand_counts.items():
    print(f"  {b}: {c}辆")

# Print all vehicles
print("\n详细列表:")
for v in vehicles:
    print(f"[{v['id']}] {v['display']} | {v['position'][:30]} | {v['powerType']} | 图:{v['photoCount']}")

# Save final JSON
output_file = os.path.join(OUTPUT_DIR, "final_vehicles.json")
with open(output_file, 'w', encoding='utf-8') as f:
    json.dump(vehicles, f, ensure_ascii=False, indent=2)
print(f"\nSaved to: {output_file}")

# Count copied images
copied = os.listdir(DEMO_IMG_DIR)
print(f"Copied {len(copied)} images to {DEMO_IMG_DIR}")
